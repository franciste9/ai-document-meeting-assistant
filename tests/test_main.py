"""Tests for the CLI entrypoint.

The summarize path uses a fake client — no network calls.
"""

import json

import docx
import pytest

from assistant import store
from assistant.errors import AssistantError
from assistant.main import ingest_file, main, summarize_document
from assistant.models import Document, SpeakerTurn

TRANSCRIPT = """\
[00:00:01] Alex: Morning everyone, let's start.
[00:00:09] Priya: I pushed the migration last night.
[00:01:14] Alex: Great, we'll review it after standup.
"""

PROSE = """\
Quarterly Review

Revenue grew faster than forecast this quarter. Hiring slowed in the second
half. Infrastructure spend was flat.
"""


@pytest.fixture
def db(tmp_path):
    return tmp_path / "test.db"


class FakeClient:
    """Records prompts and replays canned responses."""

    def __init__(self, responses=None):
        self.responses = list(responses or ['{"summary": "ok"}'])
        self.calls = []

    def complete(self, messages, system=None, **kwargs):
        self.calls.append({"messages": messages, "system": system, **kwargs})
        if len(self.responses) > 1:
            return self.responses.pop(0)
        return self.responses[0]


# --- ingest ------------------------------------------------------------------


class TestIngestFile:
    def test_ingests_plain_text(self, tmp_path, db):
        path = tmp_path / "notes.txt"
        path.write_text(PROSE, encoding="utf-8")

        doc_id, result = ingest_file(path, db_path=db)

        assert doc_id == 1
        assert result.document.doc_type == "text"
        assert "Revenue grew" in result.document.raw_text

    def test_ingests_transcript_with_speaker_turns(self, tmp_path, db):
        path = tmp_path / "meeting.txt"
        path.write_text(TRANSCRIPT, encoding="utf-8")

        _, result = ingest_file(path, db_path=db)

        assert result.document.doc_type == "transcript"
        assert len(result.document.speaker_turns) == 3
        assert result.document.speaker_turns[0].speaker == "Alex"

    def test_ingests_vtt(self, tmp_path, db):
        path = tmp_path / "meeting.vtt"
        path.write_text(
            "WEBVTT\n\n00:00:01.000 --> 00:00:04.000\n<v Alex>Morning.\n",
            encoding="utf-8",
        )

        _, result = ingest_file(path, db_path=db)

        assert result.document.doc_type == "transcript"
        assert result.document.speaker_turns[0].speaker == "Alex"

    def test_ingests_docx(self, tmp_path, db):
        path = tmp_path / "notes.docx"
        document = docx.Document()
        document.add_paragraph("First paragraph.")
        document.save(str(path))

        _, result = ingest_file(path, db_path=db)

        assert result.document.doc_type == "docx"
        assert "First paragraph." in result.document.raw_text

    def test_persists_to_the_store(self, tmp_path, db):
        path = tmp_path / "notes.txt"
        path.write_text(PROSE, encoding="utf-8")

        doc_id, _ = ingest_file(path, db_path=db)

        assert store.get_document(doc_id, db_path=db).raw_text.startswith("Quarterly")

    def test_sets_a_token_estimate(self, tmp_path, db):
        path = tmp_path / "notes.txt"
        path.write_text(PROSE, encoding="utf-8")

        _, result = ingest_file(path, db_path=db)

        assert result.document.token_estimate > 0

    def test_under_threshold_document_has_no_chunks(self, tmp_path, db):
        path = tmp_path / "notes.txt"
        path.write_text(PROSE, encoding="utf-8")

        _, result = ingest_file(path, db_path=db)

        assert result.chunks is None

    def test_over_threshold_document_is_chunked(self, tmp_path, db, monkeypatch):
        monkeypatch.setenv("CHUNK_TOKEN_THRESHOLD", "50")
        path = tmp_path / "big.txt"
        path.write_text("\n\n".join("x" * 200 for _ in range(10)), encoding="utf-8")

        _, result = ingest_file(path, db_path=db)

        assert result.chunks is not None
        assert len(result.chunks) > 1

    def test_missing_file_raises(self, tmp_path, db):
        with pytest.raises(AssistantError, match="No such file"):
            ingest_file(tmp_path / "nope.txt", db_path=db)

    def test_unsupported_type_raises(self, tmp_path, db):
        path = tmp_path / "sheet.xlsx"
        path.write_text("x", encoding="utf-8")

        with pytest.raises(AssistantError, match="Unsupported file type"):
            ingest_file(path, db_path=db)


# --- summarize ---------------------------------------------------------------


class TestSummarizeDocument:
    def _doc(self, text="short document", tokens=10):
        return Document(
            source_path="./a.txt",
            doc_type="text",
            title="A",
            raw_text=text,
            token_estimate=tokens,
        )

    def test_returns_the_model_response(self):
        client = FakeClient(['{"summary": "done"}'])
        assert summarize_document(self._doc(), client=client) == '{"summary": "done"}'

    def test_sends_the_system_prompt(self):
        client = FakeClient()
        summarize_document(self._doc(), client=client)

        assert "meeting and document analyst" in client.calls[0]["system"]

    def test_includes_document_text_in_the_user_turn(self):
        client = FakeClient()
        summarize_document(self._doc(text="the body text"), client=client)

        assert "the body text" in client.calls[0]["messages"][0]["content"]

    def test_under_threshold_makes_one_call(self):
        client = FakeClient()
        summarize_document(self._doc(tokens=10), client=client, threshold=1000)

        assert len(client.calls) == 1

    def test_over_threshold_chunks_then_merges(self):
        text = "\n\n".join(f"para{i} " + "x" * 200 for i in range(6))
        doc = self._doc(text=text, tokens=len(text) // 4)

        client = FakeClient(['{"summary": "part"}', '{"summary": "merged"}'])
        result = summarize_document(doc, client=client, threshold=100)

        # One call per chunk, plus one merge call.
        assert len(client.calls) > 2
        assert result == '{"summary": "merged"}'

    def test_merge_call_uses_the_merge_prompt(self):
        text = "\n\n".join(f"para{i} " + "x" * 200 for i in range(6))
        doc = self._doc(text=text, tokens=len(text) // 4)

        client = FakeClient(['{"summary": "part"}', '{"summary": "merged"}'])
        summarize_document(doc, client=client, threshold=100)

        assert "merging partial analyses" in client.calls[-1]["system"]


# --- CLI ---------------------------------------------------------------------


class TestCli:
    def test_ingest_prints_id_and_tokens(self, tmp_path, db, capsys):
        path = tmp_path / "notes.txt"
        path.write_text(PROSE, encoding="utf-8")

        exit_code = main(["--db", str(db), "ingest", str(path)])

        out = capsys.readouterr().out
        assert exit_code == 0
        assert "Ingested [1]" in out
        assert "token estimate:" in out

    def test_ingest_reports_speaker_turns(self, tmp_path, db, capsys):
        path = tmp_path / "meeting.txt"
        path.write_text(TRANSCRIPT, encoding="utf-8")

        main(["--db", str(db), "ingest", str(path)])

        assert "speaker turns:  3" in capsys.readouterr().out

    def test_ingest_missing_file_exits_nonzero(self, tmp_path, db, capsys):
        exit_code = main(["--db", str(db), "ingest", str(tmp_path / "nope.txt")])

        assert exit_code == 1
        assert "Error:" in capsys.readouterr().err

    def test_ingest_error_has_no_traceback(self, tmp_path, db, capsys):
        main(["--db", str(db), "ingest", str(tmp_path / "nope.txt")])

        err = capsys.readouterr().err
        assert "Traceback" not in err

    def test_list_empty_store(self, db, capsys):
        exit_code = main(["--db", str(db), "list"])

        assert exit_code == 0
        assert "No documents ingested yet." in capsys.readouterr().out

    def test_list_shows_ingested_documents(self, tmp_path, db, capsys):
        path = tmp_path / "notes.txt"
        path.write_text(PROSE, encoding="utf-8")
        main(["--db", str(db), "ingest", str(path)])
        capsys.readouterr()

        main(["--db", str(db), "list"])

        assert "[1] Quarterly Review" in capsys.readouterr().out

    def test_summarize_unknown_id_exits_nonzero(self, db, capsys):
        exit_code = main(["--db", str(db), "summarize", "--id", "999"])

        assert exit_code == 1
        assert "No document with id 999" in capsys.readouterr().err

    def test_summarize_pretty_prints_json(self, tmp_path, db, capsys, monkeypatch):
        path = tmp_path / "notes.txt"
        path.write_text(PROSE, encoding="utf-8")
        main(["--db", str(db), "ingest", str(path)])
        capsys.readouterr()

        monkeypatch.setattr(
            "assistant.main.summarize_document",
            lambda doc: '{"summary":"ok","decisions":[],"action_items":[]}',
        )
        exit_code = main(["--db", str(db), "summarize", "--id", "1"])

        out = capsys.readouterr().out
        assert exit_code == 0
        assert json.loads(out)["summary"] == "ok"

    def test_summarize_prints_non_json_response_verbatim(
        self, tmp_path, db, capsys, monkeypatch
    ):
        path = tmp_path / "notes.txt"
        path.write_text(PROSE, encoding="utf-8")
        main(["--db", str(db), "ingest", str(path)])
        capsys.readouterr()

        monkeypatch.setattr(
            "assistant.main.summarize_document", lambda doc: "not json at all"
        )
        main(["--db", str(db), "summarize", "--id", "1"])

        assert "not json at all" in capsys.readouterr().out

    def test_no_command_exits_nonzero(self, capsys):
        with pytest.raises(SystemExit) as exc:
            main([])
        assert exc.value.code != 0


class TestEndToEnd:
    def test_ingest_then_retrieve_by_id(self, tmp_path, db):
        path = tmp_path / "meeting.txt"
        path.write_text(TRANSCRIPT, encoding="utf-8")

        doc_id, _ = ingest_file(path, db_path=db)
        fetched = store.get_document(doc_id, db_path=db)

        assert fetched.doc_type == "transcript"
        assert len(fetched.speaker_turns) == 3
        assert fetched.speaker_turns[1].speaker == "Priya"

    def test_ingest_three_formats_into_one_store(self, tmp_path, db):
        txt = tmp_path / "notes.txt"
        txt.write_text(PROSE, encoding="utf-8")

        transcript = tmp_path / "meeting.txt"
        transcript.write_text(TRANSCRIPT, encoding="utf-8")

        docx_path = tmp_path / "notes.docx"
        document = docx.Document()
        document.add_paragraph("A Word document.")
        document.save(str(docx_path))

        for path in (txt, transcript, docx_path):
            ingest_file(path, db_path=db)

        stored = store.list_documents(db_path=db)
        assert len(stored) == 3
        assert {d.doc_type for d in stored} == {"text", "transcript", "docx"}
