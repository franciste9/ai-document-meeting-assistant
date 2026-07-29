"""Tests for the per-format loaders and format detection."""

import docx
import pytest
from pypdf import PdfWriter

from assistant.errors import AssistantError
from assistant.ingestion.loaders import (
    detect_format,
    load_docx,
    load_pdf,
    load_text,
    load_transcript,
)

SPEAKER_TRANSCRIPT = """\
[00:00:01] Alex: Morning everyone, let's start.
[00:00:09] Priya: I pushed the migration last night.
[00:01:14] Alex: Great, we'll review it after standup.
"""

PROSE = """\
The quarterly review covered three areas: revenue, hiring, and infrastructure.
Revenue grew faster than forecast. Hiring slowed in the second half.
Infrastructure spend was flat.
"""


# --- fixtures ----------------------------------------------------------------


@pytest.fixture
def docx_file(tmp_path):
    path = tmp_path / "notes.docx"
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.save(str(path))
    return path


@pytest.fixture
def pdf_file(tmp_path):
    """A structurally valid two-page PDF (blank pages, no text layer)."""
    path = tmp_path / "notes.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.add_blank_page(width=200, height=200)
    with path.open("wb") as handle:
        writer.write(handle)
    return path


# --- detect_format -----------------------------------------------------------


class TestDetectFormat:
    def test_pdf_by_extension(self, tmp_path):
        assert detect_format(tmp_path / "a.pdf") == "pdf"

    def test_docx_by_extension(self, tmp_path):
        assert detect_format(tmp_path / "a.docx") == "docx"

    @pytest.mark.parametrize("ext", [".vtt", ".srt"])
    def test_subtitle_formats_are_transcripts(self, tmp_path, ext):
        assert detect_format(tmp_path / f"a{ext}") == "transcript"

    def test_extension_match_is_case_insensitive(self, tmp_path):
        assert detect_format(tmp_path / "A.PDF") == "pdf"

    def test_plain_prose_txt_is_text(self, tmp_path):
        path = tmp_path / "essay.txt"
        path.write_text(PROSE, encoding="utf-8")
        assert detect_format(path) == "text"

    def test_speaker_prefixed_txt_is_transcript(self, tmp_path):
        path = tmp_path / "meeting.txt"
        path.write_text(SPEAKER_TRANSCRIPT, encoding="utf-8")
        assert detect_format(path) == "transcript"

    def test_speaker_prefixed_txt_without_timestamps_is_transcript(self, tmp_path):
        path = tmp_path / "meeting.txt"
        path.write_text(
            "Alex: Morning everyone.\nPriya: I pushed the migration.\n",
            encoding="utf-8",
        )
        assert detect_format(path) == "transcript"

    def test_empty_txt_is_text(self, tmp_path):
        path = tmp_path / "empty.txt"
        path.write_text("", encoding="utf-8")
        assert detect_format(path) == "text"

    def test_prose_with_a_colon_is_not_a_transcript(self, tmp_path):
        """A stray 'Note: ...' line shouldn't flip prose into a transcript."""
        path = tmp_path / "essay.txt"
        path.write_text(
            "Note: revenue grew this quarter.\n" + PROSE,
            encoding="utf-8",
        )
        assert detect_format(path) == "text"

    def test_unsupported_extension_raises(self, tmp_path):
        with pytest.raises(AssistantError, match="Unsupported file type"):
            detect_format(tmp_path / "sheet.xlsx")

    def test_no_extension_raises(self, tmp_path):
        with pytest.raises(AssistantError, match="Unsupported file type"):
            detect_format(tmp_path / "README")


# --- load_text ---------------------------------------------------------------


class TestLoadText:
    def test_reads_content_verbatim(self, tmp_path):
        path = tmp_path / "a.txt"
        path.write_text(PROSE, encoding="utf-8")
        assert load_text(path) == PROSE

    def test_reads_empty_file(self, tmp_path):
        path = tmp_path / "a.txt"
        path.write_text("", encoding="utf-8")
        assert load_text(path) == ""

    def test_missing_file_raises_assistant_error(self, tmp_path):
        with pytest.raises(AssistantError, match="Could not read"):
            load_text(tmp_path / "nope.txt")

    def test_undecodable_bytes_do_not_raise(self, tmp_path):
        path = tmp_path / "a.txt"
        path.write_bytes(b"caf\xff\xfe ok")
        assert "ok" in load_text(path)


# --- load_docx ---------------------------------------------------------------


class TestLoadDocx:
    def test_extracts_paragraphs(self, docx_file):
        text = load_docx(docx_file)
        assert "First paragraph." in text
        assert "Second paragraph." in text

    def test_paragraphs_are_newline_separated(self, docx_file):
        assert load_docx(docx_file) == "First paragraph.\nSecond paragraph."

    def test_corrupt_file_raises_assistant_error(self, tmp_path):
        path = tmp_path / "broken.docx"
        path.write_text("this is not a docx", encoding="utf-8")
        with pytest.raises(AssistantError, match="Could not parse DOCX"):
            load_docx(path)

    def test_missing_file_raises_assistant_error(self, tmp_path):
        with pytest.raises(AssistantError):
            load_docx(tmp_path / "nope.docx")


# --- load_pdf ----------------------------------------------------------------


class TestLoadPdf:
    def test_returns_a_string(self, pdf_file):
        assert isinstance(load_pdf(pdf_file), str)

    def test_blank_pages_yield_no_text(self, pdf_file):
        assert load_pdf(pdf_file).strip() == ""

    def test_corrupt_file_raises_assistant_error(self, tmp_path):
        path = tmp_path / "broken.pdf"
        path.write_text("this is not a pdf", encoding="utf-8")
        with pytest.raises(AssistantError, match="Could not parse PDF"):
            load_pdf(path)

    def test_missing_file_raises_assistant_error(self, tmp_path):
        with pytest.raises(AssistantError):
            load_pdf(tmp_path / "nope.pdf")


# --- load_transcript (delegates to transcripts.py) ---------------------------


class TestLoadTranscript:
    def test_returns_speaker_turns(self, tmp_path):
        path = tmp_path / "meeting.txt"
        path.write_text(SPEAKER_TRANSCRIPT, encoding="utf-8")

        turns = load_transcript(path)

        assert [t.speaker for t in turns] == ["Alex", "Priya", "Alex"]
        assert turns[0].timestamp == "00:00:01"
        assert turns[1].text == "I pushed the migration last night."

    def test_unmarked_text_falls_back_to_one_turn(self, tmp_path):
        path = tmp_path / "prose.txt"
        path.write_text(PROSE, encoding="utf-8")

        turns = load_transcript(path)

        assert len(turns) == 1
        assert turns[0].speaker is None
        assert "quarterly review" in turns[0].text

    def test_missing_file_raises_assistant_error(self, tmp_path):
        with pytest.raises(AssistantError, match="Could not read"):
            load_transcript(tmp_path / "nope.txt")
