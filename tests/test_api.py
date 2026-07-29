"""Tests for the FastAPI wrapper.

No network calls: the summarize route's dependency is overridden with a fake,
matching the pattern in test_main.py.
"""

import io
import json

import docx
import pytest
from fastapi.testclient import TestClient
from reportlab.pdfgen import canvas

from assistant import config, store
from assistant.api import app, get_max_upload_bytes, get_summarizer
from assistant.errors import AssistantError
from assistant.models import Document

TRANSCRIPT = """\
[00:00:04] Alex: Morning all. Priya, where did the migration land?
[00:00:11] Priya: It's deployed to staging. I want a day of soak before prod.
[00:00:26] Alex: Fine. Let's decide now: we ship Thursday, not Wednesday.
"""

PROSE = """\
Quarterly Review

Revenue grew faster than forecast this quarter. Hiring slowed in the second
half. Infrastructure spend was flat.
"""

SUMMARY_JSON = json.dumps(
    {
        "summary": "The team reviewed the migration and set a ship date.",
        "decisions": ["Ship Thursday rather than Wednesday."],
        "action_items": [
            {"task": "Own the prod cutover", "owner": "Priya", "due": None}
        ],
    }
)


@pytest.fixture(autouse=True)
def isolated_db(tmp_path, monkeypatch):
    """Point the store at a per-test database."""
    monkeypatch.setattr(config, "DB_PATH", str(tmp_path / "api-test.db"))
    yield


@pytest.fixture
def client():
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()


@pytest.fixture
def fake_summary(request):
    """Override the summarize dependency with a canned response."""
    payload = getattr(request, "param", SUMMARY_JSON)
    app.dependency_overrides[get_summarizer] = lambda: (lambda document: payload)
    yield payload
    app.dependency_overrides.clear()


# -- upload helpers ----------------------------------------------------------


def txt_upload(name="notes.txt", body=PROSE):
    return {"file": (name, io.BytesIO(body.encode("utf-8")), "text/plain")}


def transcript_upload(name="standup.txt"):
    return txt_upload(name, TRANSCRIPT)


def docx_upload(name="notes.docx"):
    buffer = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Q3 Planning Notes")
    document.add_paragraph("Revenue grew 12% against a 7% forecast.")
    document.save(buffer)
    buffer.seek(0)
    return {
        "file": (
            name,
            buffer,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }


def pdf_upload(name="board.pdf"):
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 800, "Board Update - Q3")
    pdf.drawString(72, 780, "Revenue finished 12 percent above plan.")
    pdf.save()
    buffer.seek(0)
    return {"file": (name, buffer, "application/pdf")}


def vtt_upload(name="review.vtt"):
    body = (
        "WEBVTT\n\n1\n00:00:02.000 --> 00:00:06.500\n"
        "<v Sam>The design review is done, we're going with option B.\n"
    )
    return {"file": (name, io.BytesIO(body.encode("utf-8")), "text/vtt")}


# -- /health -----------------------------------------------------------------


class TestHealth:
    def test_returns_200(self, client):
        assert client.get("/health").status_code == 200

    def test_reports_ok(self, client):
        assert client.get("/health").json() == {"status": "ok"}


# -- POST /documents ---------------------------------------------------------


class TestCreateDocument:
    def test_accepts_plain_text(self, client):
        response = client.post("/documents", files=txt_upload())

        assert response.status_code == 201
        assert response.json()["doc_type"] == "text"

    def test_accepts_docx(self, client):
        response = client.post("/documents", files=docx_upload())

        assert response.status_code == 201
        assert response.json()["doc_type"] == "docx"

    def test_accepts_pdf(self, client):
        response = client.post("/documents", files=pdf_upload())

        assert response.status_code == 201
        assert response.json()["doc_type"] == "pdf"

    def test_accepts_vtt_as_transcript(self, client):
        response = client.post("/documents", files=vtt_upload())

        assert response.status_code == 201
        assert response.json()["doc_type"] == "transcript"

    def test_sniffs_speaker_prefixed_txt_as_transcript(self, client):
        response = client.post("/documents", files=transcript_upload())

        body = response.json()
        assert body["doc_type"] == "transcript"
        assert body["speaker_turn_count"] == 3

    def test_returns_an_id(self, client):
        assert client.post("/documents", files=txt_upload()).json()["id"] == 1

    def test_ids_increment(self, client):
        first = client.post("/documents", files=txt_upload()).json()["id"]
        second = client.post("/documents", files=txt_upload()).json()["id"]

        assert second > first

    def test_returns_token_estimate(self, client):
        assert client.post("/documents", files=txt_upload()).json()["token_estimate"] > 0

    def test_transcript_title_uses_the_uploaded_filename(self, client):
        """The temp file's generated name must not leak into the title."""
        response = client.post("/documents", files=transcript_upload("standup.txt"))

        assert response.json()["title"] == "standup"

    def test_title_has_no_temp_file_artifacts(self, client):
        title = client.post(
            "/documents", files=transcript_upload("standup.txt")
        ).json()["title"]

        assert not title.startswith("tmp")
        assert "assistant-upload" not in title

    def test_prose_title_still_comes_from_the_body(self, client):
        """A heading in the document beats the filename."""
        response = client.post("/documents", files=txt_upload("whatever.txt"))

        assert response.json()["title"] == "Quarterly Review"

    def test_unsafe_filename_is_sanitized(self, client):
        """A path-traversal filename must not escape the temp directory."""
        response = client.post(
            "/documents", files=transcript_upload("../../../etc/passwd.txt")
        )

        assert response.status_code == 201
        assert "/" not in (response.json()["title"] or "")

    def test_filename_with_spaces_is_preserved(self, client):
        response = client.post(
            "/documents", files=transcript_upload("team standup.txt")
        )

        assert response.json()["title"] == "team standup"

    def test_prose_has_no_speaker_turn_count(self, client):
        assert client.post("/documents", files=txt_upload()).json()["speaker_turn_count"] is None

    def test_small_document_is_not_chunked(self, client):
        body = client.post("/documents", files=txt_upload()).json()

        assert body["chunked"] is False
        assert body["chunk_count"] is None

    def test_large_document_is_chunked(self, client, monkeypatch):
        monkeypatch.setenv("CHUNK_TOKEN_THRESHOLD", "50")
        big = "\n\n".join("x" * 400 for _ in range(10))

        body = client.post("/documents", files=txt_upload(body=big)).json()

        assert body["chunked"] is True
        assert body["chunk_count"] > 1

    def test_persists_to_the_store(self, client):
        doc_id = client.post("/documents", files=txt_upload()).json()["id"]

        assert store.get_document(doc_id).raw_text.startswith("Quarterly")

    def test_rejects_unsupported_type(self, client):
        files = {"file": ("sheet.xlsx", io.BytesIO(b"x"), "application/vnd.ms-excel")}

        response = client.post("/documents", files=files)

        assert response.status_code == 415
        assert "Unsupported file type" in response.json()["detail"]

    def test_missing_file_field_is_422(self, client):
        assert client.post("/documents").status_code == 422

    def test_leaves_no_temp_file_behind(self, client, tmp_path, monkeypatch):
        spool = tmp_path / "spool"
        spool.mkdir()
        monkeypatch.setattr("tempfile.tempdir", str(spool))

        client.post("/documents", files=txt_upload())

        assert list(spool.iterdir()) == []


class TestMaxUploadBytes:
    """The limit lives in api.py — it's an HTTP concern, not a CLI one."""

    def test_defaults_to_five_megabytes(self, monkeypatch):
        monkeypatch.delenv("MAX_UPLOAD_BYTES", raising=False)
        assert get_max_upload_bytes() == 5_000_000

    def test_env_var_overrides_the_default(self, monkeypatch):
        monkeypatch.setenv("MAX_UPLOAD_BYTES", "12345")
        assert get_max_upload_bytes() == 12345

    def test_blank_env_var_falls_back_to_default(self, monkeypatch):
        monkeypatch.setenv("MAX_UPLOAD_BYTES", "   ")
        assert get_max_upload_bytes() == 5_000_000

    def test_non_integer_env_var_raises(self, monkeypatch):
        monkeypatch.setenv("MAX_UPLOAD_BYTES", "five megabytes")
        with pytest.raises(AssistantError, match="must be an integer"):
            get_max_upload_bytes()

    def test_config_does_not_own_the_limit(self):
        """Regression: the constant was moved out of config.py deliberately."""
        assert not hasattr(config, "get_max_upload_bytes")
        assert not hasattr(config, "DEFAULT_MAX_UPLOAD_BYTES")


class TestUploadSizeLimit:
    def test_oversized_upload_is_413(self, client, monkeypatch):
        monkeypatch.setenv("MAX_UPLOAD_BYTES", "1000")
        oversized = "x" * 5000

        response = client.post("/documents", files=txt_upload(body=oversized))

        assert response.status_code == 413

    def test_413_explains_the_limit(self, client, monkeypatch):
        monkeypatch.setenv("MAX_UPLOAD_BYTES", "1000")

        response = client.post("/documents", files=txt_upload(body="x" * 5000))

        assert "1,000" in response.json()["detail"]

    def test_upload_at_the_limit_is_accepted(self, client, monkeypatch):
        monkeypatch.setenv("MAX_UPLOAD_BYTES", "5000")

        response = client.post("/documents", files=txt_upload(body="x" * 4000))

        assert response.status_code == 201

    def test_oversized_upload_is_not_persisted(self, client, monkeypatch):
        monkeypatch.setenv("MAX_UPLOAD_BYTES", "1000")
        client.post("/documents", files=txt_upload(body="x" * 5000))

        assert store.list_documents() == []

    def test_oversized_upload_leaves_no_temp_file(
        self, client, tmp_path, monkeypatch
    ):
        monkeypatch.setenv("MAX_UPLOAD_BYTES", "1000")
        spool = tmp_path / "spool"
        spool.mkdir()
        monkeypatch.setattr("tempfile.tempdir", str(spool))

        client.post("/documents", files=txt_upload(body="x" * 5000))

        assert list(spool.iterdir()) == []


# -- GET /documents ----------------------------------------------------------


class TestListDocuments:
    def test_empty_store_returns_empty_list(self, client):
        response = client.get("/documents")

        assert response.status_code == 200
        assert response.json() == []

    def test_lists_ingested_documents(self, client):
        client.post("/documents", files=txt_upload())
        client.post("/documents", files=transcript_upload())

        body = client.get("/documents").json()

        assert len(body) == 2

    def test_includes_ids(self, client):
        client.post("/documents", files=txt_upload())
        client.post("/documents", files=txt_upload())

        assert [d["id"] for d in client.get("/documents").json()] == [1, 2]

    def test_includes_doc_types(self, client):
        client.post("/documents", files=txt_upload())
        client.post("/documents", files=transcript_upload())

        types = [d["doc_type"] for d in client.get("/documents").json()]
        assert types == ["text", "transcript"]

    def test_does_not_include_raw_text(self, client):
        client.post("/documents", files=txt_upload())

        assert "raw_text_preview" not in client.get("/documents").json()[0]


# -- GET /documents/{id} -----------------------------------------------------


class TestGetDocument:
    def test_returns_the_document(self, client):
        doc_id = client.post("/documents", files=txt_upload()).json()["id"]

        response = client.get(f"/documents/{doc_id}")

        assert response.status_code == 200
        assert response.json()["id"] == doc_id

    def test_includes_a_text_preview(self, client):
        doc_id = client.post("/documents", files=txt_upload()).json()["id"]

        preview = client.get(f"/documents/{doc_id}").json()["raw_text_preview"]

        assert "Revenue grew" in preview

    def test_preview_is_truncated(self, client):
        doc_id = client.post(
            "/documents", files=txt_upload(body="y" * 3000)
        ).json()["id"]

        preview = client.get(f"/documents/{doc_id}").json()["raw_text_preview"]

        assert len(preview) < 600

    def test_short_text_is_not_marked_truncated(self, client):
        doc_id = client.post("/documents", files=txt_upload()).json()["id"]

        preview = client.get(f"/documents/{doc_id}").json()["raw_text_preview"]

        assert not preview.endswith("…")

    def test_includes_speaker_turn_count(self, client):
        doc_id = client.post("/documents", files=transcript_upload()).json()["id"]

        assert client.get(f"/documents/{doc_id}").json()["speaker_turn_count"] == 3

    def test_unknown_id_is_404(self, client):
        response = client.get("/documents/999")

        assert response.status_code == 404
        assert "No document with id 999" in response.json()["detail"]

    def test_non_integer_id_is_422(self, client):
        assert client.get("/documents/abc").status_code == 422

    def test_store_failure_is_502_not_404(self, client, monkeypatch):
        """A real store error must not be misreported as a missing document."""

        def boom(*args, **kwargs):
            raise AssistantError("Could not read document 1: disk failure")

        monkeypatch.setattr("assistant.api.store.get_document", boom)

        assert client.get("/documents/1").status_code == 502


# -- POST /documents/{id}/summarize ------------------------------------------


class TestSummarize:
    def test_returns_parsed_summary(self, client, fake_summary):
        doc_id = client.post("/documents", files=transcript_upload()).json()["id"]

        response = client.post(f"/documents/{doc_id}/summarize")

        assert response.status_code == 200
        assert response.json()["summary"].startswith("The team reviewed")

    def test_returns_decisions(self, client, fake_summary):
        doc_id = client.post("/documents", files=transcript_upload()).json()["id"]

        body = client.post(f"/documents/{doc_id}/summarize").json()

        assert body["decisions"] == ["Ship Thursday rather than Wednesday."]

    def test_returns_action_items(self, client, fake_summary):
        doc_id = client.post("/documents", files=transcript_upload()).json()["id"]

        items = client.post(f"/documents/{doc_id}/summarize").json()["action_items"]

        assert items[0]["owner"] == "Priya"
        assert items[0]["due"] is None

    def test_unknown_id_is_404(self, client, fake_summary):
        assert client.post("/documents/999/summarize").status_code == 404

    def test_passes_the_document_to_the_summarizer(self, client):
        seen = {}

        def fake(document):
            seen["doc"] = document
            return SUMMARY_JSON

        app.dependency_overrides[get_summarizer] = lambda: fake
        doc_id = client.post("/documents", files=transcript_upload()).json()["id"]
        client.post(f"/documents/{doc_id}/summarize")

        assert isinstance(seen["doc"], Document)
        assert seen["doc"].doc_type == "transcript"

    @pytest.mark.parametrize("fake_summary", ["not json at all"], indirect=True)
    def test_unparseable_output_is_502(self, client, fake_summary):
        doc_id = client.post("/documents", files=txt_upload()).json()["id"]

        response = client.post(f"/documents/{doc_id}/summarize")

        assert response.status_code == 502
        assert "did not return valid JSON" in response.json()["detail"]

    @pytest.mark.parametrize("fake_summary", ["not json at all"], indirect=True)
    def test_502_includes_the_raw_output(self, client, fake_summary):
        doc_id = client.post("/documents", files=txt_upload()).json()["id"]

        response = client.post(f"/documents/{doc_id}/summarize")

        assert "not json at all" in response.json()["detail"]

    @pytest.mark.parametrize(
        "fake_summary", ['{"wrong": "shape"}'], indirect=True
    )
    def test_wrong_json_shape_is_502(self, client, fake_summary):
        doc_id = client.post("/documents", files=txt_upload()).json()["id"]

        response = client.post(f"/documents/{doc_id}/summarize")

        assert response.status_code == 502
        assert "unexpected shape" in response.json()["detail"]

    def test_api_error_surfaces_as_502(self, client):
        def boom(document):
            raise AssistantError("ANTHROPIC_API_KEY is not set.")

        app.dependency_overrides[get_summarizer] = lambda: boom
        doc_id = client.post("/documents", files=txt_upload()).json()["id"]

        response = client.post(f"/documents/{doc_id}/summarize")

        assert response.status_code == 502
        assert "ANTHROPIC_API_KEY" in response.json()["detail"]


# -- OpenAPI / docs ----------------------------------------------------------


class TestDocs:
    def test_swagger_ui_is_served(self, client):
        response = client.get("/docs")

        assert response.status_code == 200
        assert "swagger" in response.text.lower()

    def test_openapi_schema_is_valid(self, client):
        schema = client.get("/openapi.json").json()

        assert schema["info"]["title"] == "Document/Meeting Assistant"

    def test_all_routes_are_documented(self, client):
        paths = client.get("/openapi.json").json()["paths"]

        assert "/health" in paths
        assert "/documents" in paths
        assert "/documents/{doc_id}" in paths
        assert "/documents/{doc_id}/summarize" in paths


class TestEndToEnd:
    def test_upload_list_fetch_summarize(self, client, fake_summary):
        created = client.post("/documents", files=transcript_upload())
        assert created.status_code == 201
        doc_id = created.json()["id"]

        listed = client.get("/documents").json()
        assert any(d["id"] == doc_id for d in listed)

        detail = client.get(f"/documents/{doc_id}").json()
        assert detail["speaker_turn_count"] == 3

        summary = client.post(f"/documents/{doc_id}/summarize").json()
        assert summary["decisions"]
